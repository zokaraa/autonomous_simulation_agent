from docx import Document
from docx.shared import Inches

def create_report():
    # Create a new Word document
    doc = Document()

    # Title
    doc.add_heading('Simulation Experiment Report: Polymer Chains', 0)

    # Abstract Section
    abstract = (
        "This report presents a detailed simulation experiment studying the properties of polymer chains in three-dimensional space. "
        "The primary objective was to analyze the end-to-end distance of polymer chains with and without self-avoidance constraints. "
        "Multiple chain lengths were considered (N = 10, 50, 100, 200), and for each length, 2000 chains were generated. "
        "Various graphs and results, including chain conformations and the mean squared end-to-end distance, were produced to understand "
        "the scaling behavior of the chains."
        "\nThe results of this simulation provide insight into the geometric and physical properties of polymer chains and can inform further research "
        "in polymer physics and material science."
    )
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(abstract)

    # Introduction Section
    introduction = (
        "Polymer chains are a crucial subject of study in materials science, playing a significant role in understanding the behavior "
        "of polymeric materials. This simulation aims to investigate the conformations of polymer chains and their properties, specifically "
        "focusing on the end-to-end distance. By employing Monte Carlo simulations, we generate large ensembles of polymer chains with "
        "random orientations in 3D space. For the self-avoiding random walk (SAW) model, an additional constraint ensures that no two segments "
        "of the chain come within a unit distance of each other. This constraint mimics the excluded volume effect observed in real polymers."
        "\nUnderstanding the behavior of such chains helps in predicting the properties of the material, such as strength, flexibility, and durability. "
        "The concept of self-avoidance is particularly critical, as it introduces realism into the model by considering physical constraints that polymers face."
        "\nThis study compares the difference between random walks and self-avoiding walks, thereby elucidating the effects of the excluded volume on the morphological "
        "properties of the chains."
    )
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(introduction)

    # Methods Section
    methods = (
        "The simulation was implemented in Python, employing various libraries such as NumPy for numerical operations and Matplotlib for "
        "visualizations. Two distinct types of polymer chains were generated: (1) random chains without self-avoidance, where each segment's "
        "orientation was uniform in 3D space; (2) self-avoiding chains, where additional checks ensured that each segment was at least one unit "
        "away from all other segments in the chain. For each value of N (10, 50, 100, 200), 2000 chains were generated. The end-to-end distance "
        "vector of each chain was computed, and the mean squared end-to-end distance h^2(N) was calculated. Scaling exponent v was determined from the "
        "relation h^2(N) ∝ N^v."
        "\nGenerating the self-avoiding chains involved a retry mechanism to ensure the placement of each new segment adheres to the self-avoidance constraint. "
        "A maximum number of attempts was set for adding each new segment, and if the chain could not be completed, it was discarded and a new chain generation was attempted."
        "\nData visualization involved plotting 50 random chains for each N value and generating summary plots of h^2(N) vs N. The results were saved as PNG images for inclusion in the report."
    )
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(methods)
    
    # Load and add figures
    fig1_path = 'Chain3D10.png'
    fig2_path = 'Chain3D50.png'
    fig3_path = 'Chain3D100.png'
    fig4_path = 'Chain3D200.png'
    fig5_path = 'h2vsN.png'
    fig6_path = 'SA_Chain3D10.png'
    fig7_path = 'SA_Chain3D50.png'
    fig8_path = 'SA_Chain3D100.png'
    fig9_path = 'SA_Chain3D200.png'
    fig10_path = 'SA_h2vsN.png'

    doc.add_heading('Results', level=1)
    results = (
        "The findings from the simulation are presented through various graphs and statistical outputs. "
        "Figures 1 through 4 depict random polymer chains without self-avoidance constraints for N=10, 50, 100, and 200 segments, respectively. "
        "Figure 5 shows the h^2(N) vs. N plot for these random chains, indicating a near-linear relationship with a scaling exponent of v ≈ 0.995."
        "\nThese plots are essential in understanding the geometric configurations that the polymer chains can adopt and show how the mean squared "
        "end-to-end distance scales with the number of segments."
    )
    doc.add_paragraph(results)
    doc.add_picture(fig1_path, width=Inches(3.5))
    doc.add_paragraph("Fig. 1: Random Polymer Chains for N=10")
    doc.add_picture(fig2_path, width=Inches(3.5))
    doc.add_paragraph("Fig. 2: Random Polymer Chains for N=50")
    doc.add_picture(fig3_path, width=Inches(3.5))
    doc.add_paragraph("Fig. 3: Random Polymer Chains for N=100")
    doc.add_picture(fig4_path, width=Inches(3.5))
    doc.add_paragraph("Fig. 4: Random Polymer Chains for N=200")
    doc.add_picture(fig5_path, width=Inches(3.5))
    doc.add_paragraph("Fig. 5: Mean Squared End-to-End Distance ($h^2(N)$) vs N for Random Chains")

    results_sa = (
        "Figures 6 through 9 illustrate self-avoiding polymer chains for N=10, 50, 100, and 200 segments, respectively. "
        "Figure 10 presents the h^2(N) vs. N plot for these self-avoiding chains, demonstrating a slightly superlinear relationship with a scaling exponent "
        "of v ≈ 1.083. These results reflect the influence of self-avoidance constraints, leading to an increase in the effective chain dimension and the end-to-end distance."
        "\nThe self-avoiding walks show a more pronounced spread in 3D space, which is indicative of the physical constraints imposed on the polymer configuration."
        "\nThus, the study provides a comprehensive contrast between the unconstrained and constrained polymer configurations, highlighting the significant effect of self-avoidance on polymer behavior."
    )
    doc.add_paragraph(results_sa)
    doc.add_picture(fig6_path, width=Inches(3.5))
    doc.add_paragraph("Fig. 6: Self-Avoiding Polymer Chains for N=10")
    doc.add_picture(fig7_path, width=Inches(3.5))
    doc.add_paragraph("Fig. 7: Self-Avoiding Polymer Chains for N=50")
    doc.add_picture(fig8_path, width=Inches(3.5))
    doc.add_paragraph("Fig. 8: Self-Avoiding Polymer Chains for N=100")
    doc.add_picture(fig9_path, width=Inches(3.5))
    doc.add_paragraph("Fig. 9: Self-Avoiding Polymer Chains for N=200")
    doc.add_picture(fig10_path, width=Inches(3.5))
    doc.add_paragraph("Fig. 10: Mean Squared End-to-End Distance ($h^2(N)$) vs N for Self-Avoiding Chains")

    # Save the document
    doc.save('Simulation_Experiment_Report.docx')

if __name__ == "__main__":
    create_report()
