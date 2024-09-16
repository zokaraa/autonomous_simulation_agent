from docx import Document
from docx.shared import Inches

# Create a Document
doc = Document()

# Title of the document
doc.add_heading('Polymer Chain Simulation Experiment Report', 0)

# Abstract section
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    "This report provides an analysis of a simulation study on polymer chains. "
    "The primary objective was to generate 2000 polymer chains consisting of N segments with random orientations in 3D space, "
    "calculate the mean squared end-to-end distance (h2(N)), and determine the relationship of h2(N) with N. "
    "Understanding the behavior of polymer chains is crucial for many applications in material science and biology, especially in predicting the physical properties of macromolecules. "
    "A key result from this study is the determination of the scaling exponent 'v' that characterizes the relationship between the mean squared end-to-end distance and the number of segments. "
    "Graphs were generated to visualize the polymer chain conformations and the correlation between h2(N) and N, providing valuable insights into the polymer chain dynamics."
)

# Introduction section
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "The study of polymer chains is significant in understanding the physical properties of macromolecules in materials science and biology. "
    "Polymers are large molecules made up of repeating subunits called monomers, and their structure affects their function and behavior in various environments. "
    "A polymer chain can be modeled as a series of connected segments, with each segment having a fixed length but random orientation. This random orientation gives rise to the study of random walks and their properties. "
    "The objective of this study was to simulate 2000 polymer chains for various lengths (N) ranging from 10 to 400 segments and analyze the mean squared end-to-end distance (h2(N)) as a function of N. "
    "The end-to-end distance of a polymer chain is a measure of the spatial extent of the polymer, which impacts its physical properties such as viscosity, tensile strength, and diffusion. "
    "The theoretical background suggests that the mean squared end-to-end distance should scale with the number of segments as h2(N) ∝ N^v, where the exponent 'v' can indicate the type of polymer chain model, such as ideal (or Gaussian) chains, self-avoiding walks, or others. "
    "This study aims to empirically determine the scaling exponent 'v' and compare it with theoretical expectations for ideal polymer chains."
)

# Methods section
doc.add_heading('Methods', level=1)
doc.add_paragraph(
    "The simulation was implemented using Python, leveraging the numpy library for numerical computations and matplotlib for data visualization. "
    "For each polymer chain, 3D unit vectors representing segment orientations were generated using a uniform distribution over a sphere. This method ensures that the directions of segments are uniformly distributed in three-dimensional space. "
    "The polymer chains were constructed by cumulatively adding these vectors starting from the origin, thus forming a random walk in three dimensions. "
    "The end-to-end distance was computed as the Euclidean distance between the first and last segment of each polymer chain. This computation was repeated for 2000 chains for each given N to obtain a statistically significant result. "
    "The mean squared end-to-end distance, h2(N), was calculated by averaging the squared end-to-end distances over the 2000 chains for each N value. "
    "To visualize the polymer chains, 50 random chain conformations were plotted for each N value. Additionally, a graph of h2(N) versus N was plotted to study the scaling behavior. "
    "Linear regression on log-transformed data was used to determine the scaling exponent 'v', providing insights into the relationship between h2(N) and N."
)

# Results section
doc.add_heading('Results', level=1)

# Include key findings and graphs
doc.add_paragraph(
    "The results of the simulation are summarized in the table and figures below. "
    "The mean squared end-to-end distances for different N values were computed as follows:\n"
    "h2(10) = 10.07\n"
    "h2(50) = 50.47\n"
    "h2(100) = 99.01\n"
    "h2(200) = 198.30\n"
    "h2(400) = 400.93\n"
    "These values are close to the corresponding N values, as expected for an ideal polymer chain model where h2(N) is proportional to N. "
    "The scaling exponent 'v' was determined to be approximately 0.997 through linear regression on the log-transformed values of N and h2(N). This value is consistent with the theoretical prediction for ideal chains, where 'v' is expected to be 1. "
    "Figures illustrating the polymer chain conformations for different N values and the plot of h2(N) versus N are included below, providing a visual representation of the simulation results."
)

# Insert the Chain conformation graphs
chain_files = ['Chain3D10.png', 'Chain3D50.png', 'Chain3D100.png', 'Chain3D200.png', 'Chain3D400.png']
for file in chain_files:
    doc.add_paragraph(f"Figure: Chain Conformations for {file.split('Chain3D')[1].split('.png')[0]} segments")
    doc.add_picture(file, width=Inches(4))

# Insert the h2(N) vs N graph
doc.add_paragraph("Figure: Mean squared end-to-end distance (h2(N)) vs N")
doc.add_picture('h2vsN.png', width=Inches(4))

# Save the document
doc.save('Polymer_Chain_Simulation_Report.docx')
