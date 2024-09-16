from docx import Document
from docx.shared import Inches

def create_report():
    document = Document()
    document.add_heading('Polymer Chain Simulation Report', level=1)

    # Abstract
    document.add_heading('Abstract', level=2)
    document.add_paragraph(
        "This report outlines the stochastic simulation of polymer chains in a 3D "
        "space, focusing on observing the scaling behavior of polymers as a function "
        "of their segment length. The primary objective is to explore the mean squared "
        "end-to-end distance of polymers with varying segment counts."
    )

    # Introduction
    document.add_heading('Introduction', level=2)
    document.add_paragraph(
        "Polymer chains are pivotal in various scientific and engineering applications. "
        "Understanding their physical behavior in space can aid in designing better materials. "
        "This study simulates polymer chains with randomly oriented segments in three-dimensional space "
        "and computes key statistical measures to analyze their structural properties."
    )

    # Methods
    document.add_heading('Methods', level=2)
    document.add_paragraph(
        "The study generated 2000 polymer chains with segment counts of 10, 50, 100, 200, and 400. "
        "Each segment's orientation was assigned randomly based upon a uniform spherical distribution. "
        "The Python programming language and libraries such as NumPy and Matplotlib were used for simulation "
        "and visualization purposes."
    )

    # Results
    document.add_heading('Results', level=2)
    document.add_paragraph(
        "The results demonstrate a linear relationship between the mean squared end-to-end distance and "
        "the number of segments in the polymer chain. ")
    document.add_paragraph("See the figures below for graphical representations of the data.")
    
    # Assuming the file paths and names are corrected here
    for N in [10, 50, 100, 200, 400]:
        image_path = f'Chain3D{N}.png'
        document.add_paragraph(f'Figure {N//10}: Polymer Chain with {N} Segments', style='Caption')
        document.add_picture(image_path, width=Inches(4.25))
        
    document.add_paragraph('Figure 5: Mean Squared End-to-End Distance vs Number of Segments', style='Caption')
    document.add_picture('h2vsN.png', width=Inches(4.25))

    document.save('Polymer_Simulation_Report.docx')

create_report()
