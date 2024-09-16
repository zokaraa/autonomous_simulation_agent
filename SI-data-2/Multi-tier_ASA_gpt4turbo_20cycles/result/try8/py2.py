from docx import Document
from docx.shared import Inches
import os

def create_report():
    # Create a new Document
    doc = Document()
    doc.add_heading('Simulation Experiment Report on Polymer Chains', level=1)

    # Adding Abstract
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(
        "This report outlines the procedures and results from a simulation study "
        "on the behavior of polymer chains in 3D space. The purpose of this study "
        "is to understand how the length of the polymer chain affects its end-to-end "
        "distance in a randomly oriented 3D space environment. Through computational "
        "simulations, we've analyzed the scaling behavior of polymer chains as their "
        "segment counts increase."
    )

    # Adding Introduction
    doc.add_heading('Introduction', level=2)
    doc.add_paragraph(
        "The physical properties of polymer chains have been a subject of extensive "
        "research due to their implications in both natural and synthetic materials. "
        "Understanding the geometric and dynamic properties of these chains in three-dimensional "
        "space is essential for developing new materials and for enhancing existing applications. "
        "In this report, we explore the end-to-end distance of polymer chains with varying segment lengths."
    )

    # Adding Methods
    doc.add_heading('Methods', level=2)
    doc.add_paragraph(
        "Our methodology involved the use of a computer simulation to generate 2000 polymer chains for "
        "each defined length segment (N = 10, 50, 100, 200, 400). Each segment of the chain was assigned "
        "a random orientation in 3D space, while ensuring a uniform distribution of the angles. The simulation "
        "was repeated multiple times to ensure statistical reliability, and for each segment count, a set of "
        "50 chain conformations were visualized and analyzed."
    )

    # Adding Results
    doc.add_heading('Results', level=2)
    doc.add_paragraph(
        "Our findings reveal a clear scaling relationship between the mean squared end-to-end distance "
        "and the number of segments in the polymer chain. The results are captured in the plotted graphs "
        "as shown below."
    )

    # Inserting images
    for N in [10, 50, 100, 200, 400]:
        filename = f'Chain3D{N}.png'
        if os.path.exists(filename):
            doc.add_paragraph(f'Fig. {N//10}: Visualization of 50 random polymer chains with N = {N}')
            doc.add_picture(filename, width=Inches(4.5))

    if os.path.exists('h2vsNGraph.png'):
        doc.add_paragraph('Fig. 5: Plot of mean squared end-to-end distance h2(N) vs. N')
        doc.add_picture('h2vsNGraph.png', width=Inches(4.5))

    # Save the document
    doc.save('Polymer_Chains_Simulation_Report.docx')
    print("Report generated successfully.")

if __name__ == "__main__":
    create_report()
