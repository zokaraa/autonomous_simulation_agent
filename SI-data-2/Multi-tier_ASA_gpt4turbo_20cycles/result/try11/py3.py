from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()
doc.add_heading('Polymer Chain Simulation Report', level=1)

# Abstract Section
doc.add_heading('Abstract', level=2)
doc.add_paragraph(
    "This report provides a detailed analysis of polymer chain simulations. "
    "The goal was to explore the end-to-end distance in polymer chains as a function of their segment length. "
    "We aimed to understand how the mean squared end-to-end distance scales with the number of segments."
)

# Introduction Section
doc.add_heading('Introduction', level=2)
doc.add_paragraph(
    "Polymer chains are crucial in various applications across multiple industries. "
    "This simulation investigates the physical properties of polymer chains as they relate to chain length. "
    "Through statistical modeling and 3D visualization, we aim to gain insights into the scaling relationships of polymers."
)

# Methods Section
doc.add_heading('Methods', level=2)
doc.add_paragraph(
    "We employed a Monte Carlo simulation approach to generate random polymer chains in three dimensions. "
    "Each chain consists of N segments, and the orientation of each segment was assigned randomly based on a uniform distribution of angles. "
    "Our analysis includes polymer chains with lengths of 10, 50, 100, 200, and 400 segments. "
    "Python and NumPy were used for computations, while Matplotlib was utilized for generating visualizations."
)

# Results Section
doc.add_heading('Results', level=2)
doc.add_paragraph(
    "The mean squared end-to-end distance was computed for each chain, and the results were plotted against chain length. "
    "Chain conformations were visualized to contribute further to our understanding. "
    "Refer to the figures below showing the scaling relationships and sample polymer chain conformations."
)

# Include the graphs with in text references
figures = [('Chain3D10.png', 10), ('Chain3D50.png', 50),
           ('Chain3D100.png', 100), ('Chain3D200.png', 200),
           ('Chain3D400.png', 400), ('h2vsN.png', 'all lengths')]

for fig_name, desc in figures:
    doc.add_paragraph(f'Fig. {fig_name[:-4]}: Visualization for N={desc}')
    doc.add_picture(fig_name, width=Inches(4.25))

# Save the document
doc.save('Polymer_Chain_Simulation_Report.docx')
print("The report has been created and saved successfully.")
