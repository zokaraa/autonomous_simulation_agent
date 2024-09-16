from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()
doc.add_heading('Polymer Chain Simulation Experiment Report', 0)

# Abstract
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    "This report outlines the results of a computational simulation designed to study the properties "
    "of polymer chains in 3D space. The objective was to explore how varying the length of polymer "
    "chains affects their end-to-end distance, hypothesizing that longer chains would exhibit a "
    "proportional increase in their squared end-to-end distance, denoted as h2(N).")

# Introduction
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "The study of polymer chains is crucial in materials science, providing insights into the "
    "macroscopic properties of materials from their microscopic structures. This report presents "
    "a simulation-based approach where polymer chains were modeled in three-dimensional space with "
    "segments assigned random orientations. We specifically looked at chains of various lengths to "
    "observe changes in their structural attributes.")

# Methods
doc.add_heading('Methods', level=1)
doc.add_paragraph(
    "We implemented a Python script to generate 2000 simulated polymer chains for specific lengths. "
    "Each segment in the chain was given a unit length and a random orientation in 3D space. The "
    "end-to-end vector from the start to the end of each polymer was calculated, followed by the "
    "mean squared end-to-end distances across all polymers for each chosen length. For visual "
    "analysis, conformation plots of 50 randomly selected chains for each length were saved.")

# Results
doc.add_heading('Results', level=1)
doc.add_paragraph(
    "Our findings indicated a direct relationship between the polymer chain length and the mean "
    "squared end-to-end distance. The graphs below visually represent conformations of various "
    "polymer lengths and the overall trend in h2(N) values as a function of N. The calculated scaling "
    "exponent closely aligns with theoretical predictions in polymer physics.")

# Add images
for N in [10, 50, 100, 200, 400]:
    doc.add_paragraph(f'Figure {N/10}: Chain3D{N}.png showing 3D conformation plots for N={N}')
    doc.add_picture(f'Chain3D{N}.png', width=Inches(3.5))
doc.add_paragraph('Figure 7: h2vsN.png showing how h2(N) varies with N')
doc.add_picture('h2vsN.png', width=Inches(4.5))

# Save the document
doc.save('Polymer_Simulation_Report.docx')
