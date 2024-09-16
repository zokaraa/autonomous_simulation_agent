from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create and title the document
doc = Document()
doc.add_heading('Polymer Chain Simulation Experiment Report', 0)

# Abstract Section
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    "This report outlines the computational simulation findings of polymer chains "
    "in 3D space. The experiment aimed to analyze the behavior of polymer chains "
    "by calculating the mean squared end-to-end distance for increasingly "
    "long polymers, examining their scaling properties relative to chain length."
)

# Introduction Section
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "The objective of this scientific experiment was to perform a series of simulations "
    "to understand the physical behavior of polymer chains as their segment length "
    "increases. Polymers are large molecules, or macromolecules, composed of many repeated subunits. "
    "Understanding their properties is crucial for applications in material science, biology, and nanotechnology."
)

# Methods Section
doc.add_heading('Methods', level=1)
doc.add_paragraph(
    "The methodology implemented involved simulating 2000 polymer chains with varying segment lengths. "
    "Each segment orientation was randomly assigned in a 3D space. The primary programming language used "
    "was Python, utilizing libraries such as NumPy for numerical operations and Matplotlib for plotting. "
    "Chains were plotted for different segment counts, and the mean squared end-to-end distance was calculated."
)

# Results Section
doc.add_heading('Results', level=1)
doc.add_paragraph(
    "The findings from the simulation experiment were significant in illustrating the effect of chain "
    "length on the end-to-end distance. The plots depicting polymer chains for N=10, 50, 100, 200, and 400 segments are shown below:"
)

# Add images
for N in [10, 50, 100, 200, 400]:
    doc.add_picture(f'Chain3D{N}.png', width=Inches(4.0))
    doc.add_paragraph(f'Fig. {N//10}: Polymer chain with {N} segments.').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph(
    "Additionally, a plot of the mean squared end-to-end distance vs. the number of segments (N) revealed "
    "a scaling relationship, indicative of a physical polymer's expansive behavior in a three-dimensional space."
)
doc.add_picture('h2vsN.png', width=Inches(5.0))
doc.add_paragraph('Fig. 6: Plot of mean squared end-to-end distance (h2) vs N.').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph(
    "The scaling exponent v, calculated from the plot, is approximately 1.03, indicating that h2(N) ≈ N^1.03. "
    "This value suggests slight deviation from ideal random walk behavior theoretically expected for polymers in a free space."
)

# Save the document
doc.save('Polymer_Simulation_Report.docx')
