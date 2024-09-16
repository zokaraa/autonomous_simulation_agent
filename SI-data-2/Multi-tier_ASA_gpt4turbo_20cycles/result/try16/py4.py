from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = Document()
doc.add_heading('Polymer Chain Analysis Report', 0)

# Abstract
doc.add_heading('Abstract', level=1)
abstract = ("This report outlines the simulation carried out to analyze the behavior of polymer chains in 3D space. "
            "The analysis focused on calculating the mean squared end-to-end distances of polymer chains and examining "
            "their scaling behavior as a function of the number of segments.")
doc.add_paragraph(abstract)

# Introduction
doc.add_heading('Introduction', level=1)
introduction = ("The objective of this experiment is to understand how the length of a polymer chain affects its "
                "end-to-end distance in a three-dimensional space. By using computational modelling, we can simulate "
                "the spatial configurations of polymer chains and obtain data on their geometrical properties.")
doc.add_paragraph(introduction)

# Methods
doc.add_heading('Methods', level=1)
methods = ("We generated multiple polymer chain samples using a computational model where each segment's orientation "
           "is assigned randomly in a 3D space. The number of segments varied among 10, 50, 100, 200, 400 for different "
           "simulations, each with 2000 chains. The mean squared end-to-end distances were calculated and plotted as "
           "a function of the chain length.")
doc.add_paragraph(methods)

# Results
doc.add_heading('Results', level=1)
results = ("The analysis shows that the mean squared end-to-end distance scales with the number of segments. "
           "The scaling exponent calculated from the simulations is approximately 1.0254, suggesting a linear "
           "relationship between the logarithm of end-to-end distance and the logarithm of the number of segments. "
           "The plots below detail these findings:")
doc.add_paragraph(results)

# Adding plots
for N in [10, 50, 100, 200, 400]:
    doc.add_paragraph(f'Figure for N={N}:')
    doc.add_picture(f'Chain3D{N}.png', width=Inches(4.5))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Plot of Mean Squared End-to-End Distance vs. Number of Segments:')
doc.add_picture('h2vsN.png', width=Inches(4.5))
last_paragraph = doc.paragraphs[-1] 
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Save the document
doc.save('Polymer_Chain_Analysis_Report.docx')

