from docx import Document
from docx.shared import Inches
import os

# Create a new Document
doc = Document()
doc.add_heading('Polymer Chain Simulation Experiment Report', level=1)

# Add Abstract section
doc.add_heading('Abstract', level=2)
doc.add_paragraph(
    "This report outlines the results of a simulation experiment conducted to analyze the behavior "
    "of polymer chains in a three-dimensional space. Each polymer chain consists of N segments of length 1, "
    "with each segment's orientation randomly assigned. The purpose of the simulation is to understand the "
    "relationship between the number of segments and the polymer's end-to-end distance in a stochastic 3D model.")

# Add Introduction section
doc.add_heading('Introduction', level=2)
doc.add_paragraph(
    "The study of polymer chains is crucial for understanding materials science, especially the properties "
    "of plastics, rubbers, and other synthetic materials. The spatial configuration of polymer chains can highly influence "
    "their macroscopic properties. The objective of this simulation is to statistically evaluate the mean squared "
    "end-to-end distance of polymer chains with varying lengths and to determine the scaling behavior as described by the exponent v.")

# Add Methods section
doc.add_heading('Methods', level=2)
doc.add_paragraph(
    "The simulation was implemented using a Python script. The ensemble of polymer chains was generated "
    "with each segment having a randomly assigned orientation in 3D space, ensuring uniform distribution over the unit sphere. "
    "For each polymer length N, 2000 chains were generated, and metrics like mean squared end-to-end distance were computed. "
    "Significant plots were generated to visually represent the data.")

# Add Results section
doc.add_heading('Results', level=2)
doc.add_paragraph(
    "The resulting mean squared end-to-end distances h2(N) were recorded for various lengths N of polymer chains. "
    "Plots for different N values (10, 50, 100, 200, 400) were saved, and a plot of h2(N) versus N was created to depict the scaling relationship. "
    "The scaling relationship exponent v was computed to be approximately 1.03. Graphical results are shown below:")
results_directory = ['.\\Chain3D10.png', '.\\Chain3D50.png', '.\\Chain3D100.png', '.\\Chain3D200.png', '.\\Chain3D400.png', '.\\h2vsN.png']
for idx, result in enumerate(results_directory, 1):
    if os.path.exists(result):
        doc.add_picture(result, width=Inches(4.5))
        doc.add_paragraph(f'Fig. {idx}. Plot from file {result}')

# Save the document
doc.save('Polymer_Chain_Simulation_Report.docx')
