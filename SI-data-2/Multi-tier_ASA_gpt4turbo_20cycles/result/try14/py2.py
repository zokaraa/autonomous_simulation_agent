from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()
doc.add_heading('Polymer Chain Simulation Report', level=1)

# Abstract
doc.add_heading('Abstract', level=2)
doc.add_paragraph(
    "This report outlines the findings from a simulation experiment designed to explore the "
    "behavior of polymer chains in a three-dimensional space. By generating a significant number "
    "of polymer chain samples, we investigated the mean squared end-to-end distance of polymers "
    "with varying segment lengths and assessed the scaling relationship between the polymer "
    "length and their spatial structure."
)

# Introduction
doc.add_heading('Introduction', level=2)
doc.add_paragraph(
    "The study of polymer chains is fundamental in understanding the macroscopic properties "
    "of polymeric materials. The conformational characteristics of polymer chains underlie "
    "many aspects of material science and biophysics. This report examines the end-to-end "
    "distance of polymer chains as a function of their segment count through computational "
    "simulation, to establish a scaling law that describes their expanding behavior as they grow in length."
)

# Methods
doc.add_heading('Methods', : Receipt of Delivery Receipt of COVID-19 Support Solutions Reportevel=2)
doc.add_paragraph(
    "We simulated 2000 independent polymer chains, each consisting of varying segments (N = 10, 50, 100, 200, 400). "
    "Each segment was assigned a random orientation in three-dimensional space according to a uniform "
    "distribution on a unit sphere. Polymer chains were then constructed by adding segments end-to-end, and "
    "the end-to-end distance for each chain was computed. Graphical plots of chain conformations and end-to-end distances "
    "were generated without visualization to ensure computation-centric evaluation."
)

# Results
doc.add_heading('Results', level=2)
doc.add_paragraph(
    "The mean squared end-to-end distances were found to scale with the number of segments in a polymer chain. "
    "The following figures illustrate the conformation of chains and how mean squared distances vary with the number of segments:"
)

# Insert the plots
for N in [10, 50, 100, 200, 400]:
    doc.add_paragraph(f'Fig. {int(N/10)} - 3D Conformations for N={N}')
    doc.add_picture(f'Chain3D{N}.png', width=Inches(4.0))

doc.add_paragraph('Fig. 6 - Scaling of Mean Squared End-to-End Distance with Chain Length')
doc.add_picture('h2vsN.png', width=Inches(4.0))

doc.add_paragraph(
    "The analysis confirmed a scaling relationship between the mean squared end-to-end distance h2(N) and the number "
    "of segments N, substantiated by a scaling exponent derived as v = 1.039. This value suggests a nearly ideal random "
    "coil behavior in three-dimensional space."
)

# Save the document
doc.save('Polymer_Chain_Simulation_Report.docx')
