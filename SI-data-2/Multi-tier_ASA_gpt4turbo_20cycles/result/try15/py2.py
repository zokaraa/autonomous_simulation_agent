from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text):
    """Utility function to add and format heading."""
    par = doc.add_heading(text, level=1)
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_paragraph(doc, text, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """Utility function to add and format paragraph."""
    par = doc.add_paragraph(text)
    par.alignment = alignment
    for run in par.runs:
        run.font.size = Pt(12)

def create_doc():
    doc = Document()
    doc.add_heading('Polymer Chain Simulation Report', 0)

    # Abstract
    add_heading(doc, 'Abstract')
    add_paragraph(doc, ("This document presents the findings from a simulation experiment involving "
                        "the generation of 2000 polymer chains with varied segment lengths. The purpose of this "
                        "simulation was to analyze the behavior of the polymer chains in 3D space, focusing on their "
                        "mean squared end-to-end distance."))

    # Introduction
    add_heading(doc, 'Introduction')
    add_paragraph(doc, ("The study of polymer chains in three-dimensional space provides valuable insights into material science, "
                        "chemistry, and physics. By simulating different configurations of polymer chains, researchers can understand "
                        "how the length of the chain affects its overall properties."))

    # Methods
    add_heading(doc, 'Methods')
    add_paragraph(doc, ("The polymer chains were each comprised of N segments with a length of 1. The orientation of each "
                        "segment was assigned randomly in 3D space, ensuring a uniform distribution. For our analysis, "
                        "we considered chain lengths of N = 10, 50, 100, 200, and 400."))

    # Results
    add_heading(doc, 'Results')
    add_paragraph(doc, ("Our findings show a clear pattern of the mean squared end-to-end distance varying with the number of segments. "
                        "Specifically, we observed the scaling relationship known as 'Flory's idealized model' in polymer physics "
                        "(Fig. 1 - h2vsN.png). Additionally, selected conformations of polymer chains for each N have been visualized and "
                        "collated into respective figures (e.g., Fig. 2 - Chain3D10.png, Fig. 3 - Chain3D50.png)."))

    doc.save('Simulation_Experiment_Report.docx')

if __name__ == "__main__":
    create_doc()
