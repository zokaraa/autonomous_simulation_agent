from docx import Document
from docx.shared import Inches

def create_report():
    # Create a new Word document
    doc = Document()
    doc.add_heading('Polymer Chain Simulation Experiment Report', level=1)

    # Abstract section
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(
        "This report outlines the results of a simulation performed to "
        "analyze the behavior of polymer chains. Each polymer chain consists "
        "of segments that randomly orient in a 3D space. The main focus was on "
        "determining the mean squared end-to-end distance for varying numbers "
        "of segments and understanding how this measure scales with the number "
        "of segments."
    )

    # Introduction section
    doc.add_heading('Introduction', level=2)
    doc.add_paragraph(
        "The study of polymer chains is significant in materials science where "
        "understanding the physical properties of polymer materials is crucial. "
        "This simulation provides insights into the scaling properties of polymers, "
        "which is essential for the development of new materials."
    )

    # Methods section
    doc.add_heading('Methods', level=2)
    doc.add_paragraph(
        "The simulation was conducted by generating 2000 instances of polymer chains "
        "with segments ranging from 10 to 400. Each segment's orientation was randomly "
        "assigned in three-dimensional space to model realistic physical properties "
        "of polymers. The mean squared end-to-end distance was computed for each set of "
        "polymer chains, analyzing the relationship between chain length and spatial characteristics."
    )

    # Results section
    doc.add_heading('Results', level=2)
    doc.add_paragraph(
        "The results demonstrate a clear scaling relationship between the polymer length "
        "(number of segments) and the mean squared end-to-end distance. The collected data "
        "suggests that the mean squared distance scales with the number of segments according "
        "to a power law relationship. Graphs representing the simulations and computed metrics "
        "are included below."
    )

    # Inserting graphs
    for N in [10, 50, 100, 200, 400]:
        doc.add_picture(f'Chain3D{N}.png', width=Inches(4.25))
        doc.add_paragraph(f'Fig. {int(N/10)}: Polymer chain simulations with {N} segments.')

    doc.add_picture('h2vsN.png', width=Inches(4.25))
    doc.add_paragraph('Fig. 6: Plot of the mean squared end-to-end distance vs. number of segments.')

    # Save the document
    doc.save('Polymer_Simulation_Report.docx')

if __name__ == "__main__":
    create_report()
