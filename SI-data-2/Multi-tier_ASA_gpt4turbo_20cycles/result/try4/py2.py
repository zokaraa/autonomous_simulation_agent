from docx import Document
from docx.shared import Inches

def create_document():
    doc = Document()
    doc.add_heading('Polymer Chain Simulation Report', level=1)

    # Abstract
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(
        "This document outlines the simulation experiment conducted to examine "
        "the behavior of polymer chains in 3D space. The purpose of this simulation "
        "is to understand how the end-to-end distance of polymer chains varies as a function "
        "of the number of segments, applying principles of molecular dynamics and statistical mechanics.")

    # Introduction
    doc.add_heading('Introduction', level=2)
    doc.add_paragraph(
        "Polymers are large molecules composed of repeating structural units, and their molecular behavior "
        "is of significant interest in materials science. This report details the simulations carried out "
        "to model polymer chains with varying lengths and analyze their geometric properties in three dimensions.")

    # Methods
    docorpmentent('Methods', level=2)
    doc.add_paragraph(
        "We used Python for simulation, specifically utilizing numpy for numerical operations and matplotlib "
        "for plotting. Polymer chains were simulated as random walks in three-dimensional space with each segment "
        "assigned a random orientation uniformly distributed across all possible directions.")

    # Results
    doc.add_heading('Results', level=2)
    doc.add_paragraph(
        "The results from the simulations have been compiled and analyzed. Figures display the polymer chains "
        "from different simulations and their corresponding mean squared end-to-end distances plotted against "
        "the number of segments.")
    for N in [10, 50, 100, 200, 400]:
        doc.add_paragraph(f'Figure {N//10}: Simulation of 50 polymer chains with {N} segments each.')
        doc.add_picture(f'Chain3D{N}.png', width=Inches(5.0))
    
    doc.add_paragraph("Figure 5: Plot of mean squared end-to-end distance vs. number of segments")
    doc.add_picture('h2vsN.png', width=Inches(5.0))

    # Save the document
    doc.save('Polymer_Simulation_Report.docx')

if __name__ == '__main__':
    create_document()
