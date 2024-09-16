from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    return paragraph

def create_report():
    doc = Document()
    doc.add_heading('Simulation Experiment Report', 0)

    # Abstract
    add_heading(doc, 'Abstract', level=1)
    abstract_text = ("This report outlines the simulation results of generating "
                     "2000 random polymer chains with varying segment lengths and "
                     "evaluating their structural properties in 3-D space.")
    add_paragraph(doc, abstract, ext)

    # Introduction
    add_heading(doc, 'Introduction', level=1)
    intro_text = ("The purpose of this analysis is to explore the behavior of "
                  "polymer chains by simulating their random walks in a three-dimensional space. "
                  "The focus is on determining the scaling relationship of mean squared end-to-end distance "
                  "with respect to the number of segments in a polymer chain.")
    add_paragraph(doc, intro_text)

    # Methods
    add_heading(doc, 'Methods', level=1)
    methods_text = ("We employed a Python script to generate chains wherein each segment's orientation "
                    "was assigned randomly in 3D space with uniform distribution. Specifically, segment angles "
                    "were chosen to ensure a statistically uniform distribution over the sphere. "
                    "The end-to-end distance vector was computed for each polymer chain, and the mean squared "
                    "distance was calculated by averaging over 2000 samples.")
    add_paragraph(doc, methods_text)

    # Results
    add_heading(doc, 'Results', level=1)
    results_text = ("The mean squared end-to-end distance h2(N) was plotted against the number of segments N, "
                    "demonstrating a clear scaling relationship. The scaling exponent v, estimated through linear regression "
                    "of the log-transformed values, was found to be approximately 1.022. "
                    "Simulated chain conformations are shown in Fig. 1 to Fig. 5, corresponding to different segment counts.")
    results_paragraph = add_paragraph(doc, results_text)
    for i, N in enumerate([10, 50, 100, 200, 400], 1):
        doc.add_picture(f'Chain3D{N}.png')
        results_paragraph.add_run(f'\nFig. {i}: 3D view of 50 random polymer chain conformations with {N} segments.\n')

    doc.save('Simulation_Experiment_Report.docx')

if __name__ == "__main__":
    create_report()
