from docx import Document
from docx.shared import Inches
import os

def add_image_to_document(doc, image_path, caption):
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = 1  # center align image
        doc.add_paragraph(f"Figure: {caption}", style='Caption')

def create_report():
    doc = Document()
    doc.add_heading('Simulation Experiment Report on Polymer Chains', level=1)

    # Abstract
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(
        "This report summarises the findings from a simulation study "
        "of polymer chains, focusing on their three-dimensional configurations "
        "and the scaling of end-to-end distances. The analysis employs computational "
        "simulations to explore the behavior of polymer chains with varying segment lengths."
    )

    # Introduction
    doc.add_heading('Introduction', level=2)
    doc.add_paragraph(
        "Polymer chain analysis in three-dimensional space is crucial for understanding "
        "the microscopic properties of polymers. This simulation experiment aimed to "
        "investigate the mean squared end-to-end distance as a function of segment number, "
        "illustrating fundamental properties of polymers in solution."
    )

    # Methods
    doc.add_heading('Methods', level=2)
    doc.add_paragraph(
        "Using a Python script, 2000 polymer chains for each specified segment length (N=10, 50, 100, 200, 400) were generated. "
        "Each segment of these chains was randomly oriented in 3D space. The end-to-end distance was calculated for each chain, "
        "and these values were used to derive the mean squared end-to-end distances and to analyze scaling behaviors."
    )

    # Results
    doc.add_heading('Results', level=2)
    doc.add_paragraph(
        "The simulation findings reveal a consistent scaling relationship demonstrated through the mean squared end-to-end distance. "
        "The plots below illustrate the polymer configurations for different segment numbers and the logarithmic plot of mean squared end-to-end distance against segment number."
    )
    for N in [10, 50, 100, 200, 400]:
        image_path = f'Chain3D{N}.png'
        add_image_to_document(doc, image_path, f'Chain configurations with N={N} segments.')

    add_image_to_document(doc, 'h2vsN.png', 'Log-log plot of mean squared end-to-end distance vs. segment number.')
    
    # Saving document
    doc.add_page_break()
    doc.save('Polymer_Chain_Analysis_Report.docx')

if __name__ == "__main__":
    create_report()
