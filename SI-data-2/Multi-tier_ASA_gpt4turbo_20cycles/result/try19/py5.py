from docx import Document
from docx.shared import Inches

def create_report():
    # Create a new Document
    doc = Document()
    doc.add_heading('Polymer Chain Simulation Report', level=1)

    # Abstract Section
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(
        "This report presents the findings from a computational experiment "
        "designed to simulate and analyze the behavior of 3D polymer chains. "
        "Each chain consists of segments with randomized orientations, and the "
        "aim was to determine the mean squared end-to-end distances and explore "
        "its scaling behaviors."
    )

    # Introduction Section
    doc.add_heading('Introduction', level=2)
    doc.add_paragraph(
        "The study of polymer chains in three-dimensional space provides "
        "insight into the physical properties of materials at the molecular level. "
        "This simulation addresses the fundamental characteristics of polymer chains, "
        "particularly focusing on the statistical property of mean squared end-to-end distance."
    )

    # Methods Section
    doc.add_heading('Methods', level=2)
    doc.add_paragraph(
        "Using Python, 2000 simulations were performed for polymer chains with varying lengths. "
        "Each segment of the chain was assigned a random three-dimensional orientation. "
        "The simulation focused on five different chain lengths (N=10, 50, 100, 200, 400), "
        "computing the end-to-end distance vector for each polymer."
    )

    # Results Section
    doc.add_heading('Results', level=2)
    doc.add_paragraph(
        "The mean squared end-to-ending distance was determined for each set and the scaling relationship "
        "was analyzed. The following figures illustrate various polymer chain configurations and "
        "the corresponding end-to-end distance statistics."
    )
    for N in [10, 50, 100, 200, 400]:
        doc.add_picture(f'Chain3D{N}.png', width=Inches(4.25))
        doc.add_paragraph(f"Fig. {N // 10 - 1}: Polymer Chain Simulation with N={N}")

    doc.add_picture('h2vsN.png', width=Inches(4.25))
    doc.add_paragraph("Fig. 5: Log-Log Plot of Mean Squared End-to-End Distance vs. Polymer Length")

    # Save the document
    doc.save('Polymer_Simulation_Report.docx')

if __name__ == "__main__":
    create_report()
