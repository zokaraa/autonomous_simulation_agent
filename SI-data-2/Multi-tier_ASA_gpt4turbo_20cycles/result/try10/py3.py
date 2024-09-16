from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document
doc = Document()
doc.add_heading('Simulation Experiment Report on Polymer Chains', level=1)

# Abstract
doc.add_heading('Abstract', level=2)
abstract_text = (
    "This report outlines a simulation experiment designed to investigate the "
    "behavior of polymer chains in 3D space. The experiment calculates the mean squared "
    "end-to-end distance of polymer chains for various segment lengths and examines "
    "the scaling behavior of these distances as a function of the number of segments."
)
doc.add_paragraph(abstract_text)

# Introduction
doc.add_heading('Introduction', level=2)
introduction_text = (
    "Polymer chain modeling is pivotal in understanding material properties at "
    "the molecular level. In this experiment, we simulate the random configuration "
    "of polymer chains in a three-dimensional space. The objective is to explore "
    "how the chain length affects the spatial configuration, quantified by the end-to-end distance."
)
doc.add_paragraph(introduction_text)

# Methods
doc.add_heading('Methods', level=2)
methods_text = (
    "We implemented the simulation using Python. This involved generating unit vectors "
    "representing polymer segments randomly oriented in 3D space. For each polymer chain, "
    "consisting of segments numbering 10, 50, 100, 200, and 400, we calculated the mean "
    "squared end-to-end distance using numpy for vector arithmetic and matplotlib for visual representation. "
)
doc.add_paragraph(methods_text)

# Results
doc.add_heading('Results', level=2)
results_text = (
    "The results reveal that the mean squared end-to-end distance increases with "
    "the number of segments, following a scaling relation. Plots of these polymer "
    "chain conformations and their corresponding end-to-end distances can be found in "
    "Figures 1 through 6."
)
doc.add_paragraph(results_text)

# Add image references
filenames = ["Chain3D10.png", "Chain3D50.png", "Chain3D100.png", "Chain3D200.png", "Chain3D400.png", "h2vsN.png"]
for i, filename in enumerate(filenames, 1):
    doc.add_paragraph(f'Figure {i}: {filename}', style='Caption')
    doc.add_picture(filename, width=Inches(4.25))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Save the document
doc.save('Polymer_Simulation_Report.docx')
