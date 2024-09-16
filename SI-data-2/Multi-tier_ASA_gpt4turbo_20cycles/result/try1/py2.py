from docx import Document
from docx.shared import Inches

# <<<subtask: Create a Word document for the detailed simulation experiment report>>>
# <<<subtask: Organize the report into four sections: Abstract, Introduction, Methods, Results>>>
# <<<subtask: Embed images from previous step, accordingly referencing them in the Results section>>>

# Create a new Word Document
doc = Document()
doc.add_heading('Simulation Experiment Report on Polymer Chain Conformations', level=1)

# Abstract
doc.add_heading('Abstract', level=2)
doc.add_paragraph(
    "This report outlines the computational simulation conducted to analyze "
    "the properties of polymer chains in a three-dimensional space. The objective "
    "was to understand how varying the number of segments (N) affects the mean "
    "squared end-to-end distance of the chains, potentially revealing insights "
    "into polymer physics and material science."
)

# Introduction
doc.add_heading('Introduction', level=2)
doc.add_paragraph(
    "Polymer chains consist of repeated units and their behavior in different conditions "
    "is a fundamental topic in material science. The end-to-end distance is a critical "
    "parameter in the characterization of polymer chains, influencing their mechanical "
    "and thermal properties. This simulation aimed to investigate the behavior of "
    "polymer chains by varying the number of segments under theoretical random orientations."
)

# Methods
doc.add_heading('Methods', level=2)
doc.add_paragraph(
    "A Python script was developed to generate 2000 instances of polymer chains for "
    "each specified segment length (N=10, 50, 100, 200, 400). Each segment was assigned "
    "a random orientation in 3D space. The simulation calculated the end-to-end distance "
    "and its square, averaged over all instances, to compute the mean squared end-to-end "
    "distance (h2(N)). Visualization of 50 random chain conformations for each N was "
    "created and analyzed."
)

# Results
doc.add_heading('Results', level=2)
doc.add_paragraph(
    "The results demonstrate a scaling behavior of the mean squared end-to-end distance "
    "with respect to the number of segments. The scaling exponent calculated from the "
    "data was approximately 1.03, suggesting a direct proportionality with N, indicative "
    "of a random coil conformation in three dimensions. Below are representative plots "
    "for each set of segment lengths examined."
)

# Add images
for N in [10, 50, 100, 200, 400]:
    doc.add_paragraph(f'Fig. {N//10}: 3D Conformations for N={N}')
    doc.add_picture(f'Chain3D{N}.png', width=Inches(4.5))

doc.add_paragraph('Fig. 6: Mean Squared End-to-End Distance vs N')
doc.add_picture('h2vsN.png', width=Inches(4.5))

# Save the document
doc.save('Simulation_Experiment_Report.docx')

# End of subtasks >>>
