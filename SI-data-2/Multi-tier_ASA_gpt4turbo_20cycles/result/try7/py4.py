from docx import Document
from docx.shared import Inches

def create_report():
    # Create a new Document
    doc = Document()
    doc.add_heading('Polymer Chain Simulation Experiment', level=1)

    # Abstract
    doc.add_heading('Abstract', level=2)
    abstract = (
        "This report summarizes the outcomes of a simulation experiment conducted to "
        "analyze the properties of randomly oriented polymer chains in 3D space. "
        "Specifically, the focus was on calculating the mean squared end-to-end distance "
        "of the chains and exploring the scaling behaviors as the number of segments increases."
    )
    doc.add_paragraph(abstract)

    # Introduction
    doc.add_heading('Introduction', level=2)
    introduction = (
        "Polymer chains can exhibit complex behaviors depending on their molecular structure "
        "and the nature of their constituent segments. Understanding the spatial configuration "
        "of these polymers through computational simulations helps in predicting their physical "
        "properties and usefulness in various applications. This experiment seeks to provide insights "
        "into the behavior of polymer chains through a computational exploration."
    )
    doc.add_paragraph(introduction)

    # Methods
    doc.add_heading('Methods', level=2)
    methods = (
        "We simulated 2000 polymer chains with a fixed segment length but varying total lengths consisting "
        "of 10, 50, 100, 200, and 400 segments. Each segment's orientation was determined randomly in 3D space, "
        "ensuring a uniform distribution of angles. The simulation calculates the end-to-end distance vector "
        "and the mean squared end-to-end distance for each set of polymer chains."
    )
    doc.add_paragraph(methods)

    # Results
    doc.add_heading('Results', level=2)
    results = (
        "The distributions and behaviors of the polymer chains were visualized through multiple plots, "
        "highlighting the variances and patterns as the number of segments increased. "
        'Key figures include the plots of the polymer chains ("Fig. 1-5") and the plot of mean squared '
        'end-to-end distance against the number of segments ("Fig. 6"). The observed scaling exponent '
        f'was found to be approximately {1.0051124288768594}, indicating a nearly linear relationship '
        'between the number of segments and the mean squared end-to-end distance.'
    )
    doc.add_paragraph(results)
    doc.add_picture('h2vsN.png', width=Inches(6))

    for N in [10, 50, 100, 200, 400]:
        doc.add_picture(f'Chain3D{N}.png', width=Inches(6))

    # Save the report
    doc.save('Polymer_Chain_Simulation_Report.docx')

if __name__ == "__main__":
    create_report()
