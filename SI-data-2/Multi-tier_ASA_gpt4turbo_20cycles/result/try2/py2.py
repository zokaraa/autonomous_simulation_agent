from docx import Document
from docx.shared import Inches
import os

def add_image(document, image_title, image_path):
    document.add_heading(image_title, level=3)
    document.add_picture(image_path, width=Inches(4.0))
    document.add_paragraph("Figure: " + image_title)

# Create a new Word document
doc = Document()
doc.add_heading('Polymer Chain Simulation Report', level=1)

# Abstract Section
doc.add_heading('Abstract', level=2)
abstract_text = (
    "This document outlines the results of a simulation experiment conducted to analyze "
    "the properties of polymer chains. Each polymer was modeled as a random walk in three "
    "dimensions, with chains consisting of varying lengths. The purpose of the experiment "
    "was to understand how the mean squared end-to-end distance of polymer chains scales "
    "with the number of segments in the chain."
)
doc.add_paragraph(abstract_text)

# Introduction Section
doc.add_heading('Introduction', level=2)
introduction_text = (
    "The experiment aimed to investigate the behavior of polymer chains in 3D space. "
    "Polymer chains are essential in various applications, including materials science "
    "and biophysics, making understanding their structural properties critical. "
    "The focus was on how the chain length affects the spatial configuration, "
    "represented by the mean squared end-to-end distance."
)
doc.add_paragraph(introduction_text)

# Methods Section
doc.add_heading('Methods', level=2)
methods_text = (
    "We simulated 2000 polymer chains for lengths of 10, 50, 100, 200, and 400 segments. "
    "Each segment was randomly oriented to simulate a true random walk in 3D space. "
    "The simulation code was written in Python, utilizing libraries such as NumPy for "
    "mathematical operations, and Matplotlib for visualization. "
    "The mean squared end-to-end distances were computed and analyzed. "
)
doc.add_paragraph(methods_text)

# Results Section
doc.add_heading('Results', level=2)
results_text = (
    "The experiments provided insights into chain conformation as a function of chain length. "
    f"The calculated scaling exponent of the mean squared end-to-end distance was approximately 0.986. "
    "This value indicates a nearly linear relationship between chain length and spatial extension in the conditions explored. "
)
doc.add_paragraph(results_text)

# Add all necessary images from the task
image_files = [file for file in os.listdir() if file.startswith('Chain3D') or file == 'h2vsN.png']
for image_file in image_files:
    add_image(doc, image_file.replace('.png', ''), image_file)

# Save the document
doc.save('Polymer_Chain_Simulation_Report.docx')
