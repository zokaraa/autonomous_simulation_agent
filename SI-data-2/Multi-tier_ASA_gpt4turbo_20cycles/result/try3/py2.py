from docx import Document
from docx.shared import Inches

def create_document():
    doc = Document()
    
    doc.add_heading('Simulation Experiment Report on Polymer Chains', level=1)

    # Abstract
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(
        "This report outlines the computational simulation performed to analyze the behavior of polymer "
        "chains in 3D. The experiment entailed generating multiple configurations of polymer chains, each "
        "consisting of varying segment lengths, and analyzing their mean squared end-to-end distances. This analysis "
        "helps in understanding the scaling behavior of polymers in spatial configurations."
    )

    # Introduction
    doc.add_heading('Introduction', level=2)
    doc.add_paragraph(
        "Polymers are substantial macromolecules composed of repeating structural units. The physical properties "
        "of polymers depend heavily on their shapes and structures in three-dimensional space. Therefore, it becomes "
        "important to simulate and understand these structures for material engineering and science. This report "
        "discusses a computational approach to simulate and visualize polymer chains in a three-dimensional space and "
        "investigate the relationship between chain length and end-to-end distance."
    )

    # Methods
    doc.add_heading('Methods', level=2)
    doc.add_paragraph(
        "Each polymer chain consisted of N segments, each of unit length, randomly oriented in 3D space. "
        "The orientations were uniformly distributed, calculated using spherical coordinates transformation. A total of "
        "2000 chains were generated for each specified length of N segments. We computed the mean squared end-to-end "
        "distance for these chains and plotted them for visual analysis and further understanding. "
        "The simulations were performed using Python, leveraging libraries like NumPy for mathematical operations "
        "and Matplotlib for plotting."
    )

    # Results
    doc.add_heading('Results', level=2)
    doc.add_paragraph(
        "The results of the simulation are visually represented as plots for each specified N (segments count). "
        "These graphs demonstrate how the chain segments assume varied configurations in a three-dimensional space. "
        "Moreover, a plot of mean squared end-to-end distance versus number of segments is generated to investigate "
        "the scaling behavior."
    )

    image_files = {
        'Chain3D10.png': 'Fig. 1: 3D Visualization for N=10',
        'Chain3D50.png': 'Fig. 2: 3D Visualization for N=50',
        'Chain3D100.png': 'Fig. 3: 3D Visualization for N=100',
        'Chain3D200.png': 'Fig. 4: 3D Visualization for N=200',
        'Chain3D400.png': 'Fig. 5: 3D Visualization for N=400',
        'h2vsN.png': 'Fig. 6: h2(N) vs N Plot on a Log-Log scale'
    }
    
    for img_path, caption in image_files.items():
        doc.add_picture(img_path, width=Inches(6.0))
        doc.add_paragraph(caption)

    doc.save('Polymer_Simulation_Report.docx')

if __name__ == "__main__":
    create_document()
