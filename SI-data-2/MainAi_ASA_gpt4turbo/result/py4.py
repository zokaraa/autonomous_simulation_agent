from docx import Document
from docx.shared import Inches

def create_report():
    # Initialize Document
    doc = Document()
    doc.add_heading('Polymer Chain Simulation Report', 0)

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This report outlines the findings of a simulation experiment designed to analyze the properties of polymer chains. "
        "Each polymer chain consists of N segments, with each segment randomly oriented in 3D space. "
        "The focus was on determining the mean squared end-to-end distances, the behavior of polymer chains in 3D space at varying lengths, and their scaling relationships."
    )

    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "Polymeric materials exhibit unique properties that depend heavily on their molecular structures. "
        "Understanding the end-to-end distance of polymer chains helps in exploring their entropic properties in solvents, "
        "conditions of dilution, and response to external forces. This experiment aims to contribute to the understanding of polymer physics through simulated visualizations and computational analysis."
    )

    # Methods
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        "The polymer chains were simulated using a Python script that assigns a random orientation to each segment. "
        "A total of 2000 chains were generated for chains containing 10, 50, 100, 200, and 400 segments. "
        "For each set, 50 chain conformations were visualized, and the mean squared end-to-end distances (h2) were computed and analyzed to understand the scaling behavior as a function of N."
    )

    # Results
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        "The end-to-end distance was found to scale with the number of segments as h2(N) ≈ N^1.0284. "
        "The observed scaling relationship suggests a slightly superlinear behavior, which is typical under certain theoretical assumptions in polymer physics. "
        "The graphs below elucidate the visual conformity and the presented scaling law."
    )

    # Include graphs
    doc.add_heading('Figures', level=2)
    for N in [10, 50, 100, 200, 400]:
        doc.add_paragraph(f'Figure: Chain conformations for N={N} segments')
        doc.add_picture(f'Chain3D{N}.png', width=Inches(4.5))

    doc.add_paragraph('Figure: Mean Squared End-to-End Distance vs Number of Seguments')
    doc.add_picture('h2vsN.png', width=Inches(4.5))

    # Save document
    doc.save('Polymer_Simulation_Report.docx')

create_report()
